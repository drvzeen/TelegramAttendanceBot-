import logging
from datetime import datetime
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
import math
import os

# ================== НАСТРОЙКИ ==================
# Берём токен из переменных окружения. Это безопасный способ.
TOKEN = os.getenv("TOKEN") 
# Координаты центра университета для проверки местоположения.
UNIVERSITY_CENTER = (41.351376, 69.221844)  
# Разрешённый радиус (в метрах) для отметки присутствия.
ALLOWED_RADIUS = 100  
# ===============================================

# USERS: username -> {"name": ФИО, "role": "student"/"leader"}
# Внимание: эти данные будут сброшены при каждом перезапуске бота.
# Для постоянного хранения нужен файл или база данных.
USERS = {}  

# Словарь для хранения отметок посещаемости.
attendance = {}  

logging.basicConfig(level=logging.INFO)

# ================== ФУНКЦИИ-ПОМОЩНИКИ ==================
def distance(coord1, coord2):
    """
    Вычисляет расстояние между двумя координатами (широта, долгота) в метрах,
    используя формулу Гаверсина.
    """
    R = 6371000 # Радиус Земли в метрах
    lat1, lon1 = map(math.radians, coord1)
    lat2, lon2 = map(math.radians, coord2)
    dlat = lat2 - lat1
    dlon = lon2 - lon1
    a = math.sin(dlat / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2) ** 2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    return R * c

def is_leader(username):
    """Проверяет, является ли пользователь лидером."""
    return username in USERS and USERS[username]["role"] == "leader"

def is_student(username):
    """Проверяет, является ли пользователь студентом."""
    return username in USERS and USERS[username]["role"] == "student"

import json

USERS_FILE = "users.json"
ATTENDANCE_FILE = "attendance.json"

def load_data():
    """Загружает данные пользователей и посещаемости из JSON-файлов."""
    global USERS, attendance
    try:
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE, "r") as f:
                USERS = json.load(f)
        if os.path.exists(ATTENDANCE_FILE):
            with open(ATTENDANCE_FILE, "r") as f:
                attendance = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        # Если файлы не найдены или повреждены, просто начнём с пустых данных.
        pass

def save_data():
    """Сохраняет данные пользователей и посещаемости в JSON-файлы."""
    with open(USERS_FILE, "w") as f:
        json.dump(USERS, f, indent=4)
    with open(ATTENDANCE_FILE, "w") as f:
        json.dump(attendance, f, indent=4)

# ================== КОМАНДЫ ДЛЯ БОТА ==================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отправляет приветственное сообщение при вызове /start."""
    user = update.message.from_user
    await update.message.reply_text(
        f"Привет, {user.first_name}! 👋\n"
        f"Отправь '+' если ты на паре, '-' если нет.\n"
        f"Можно отправить геолокацию для подтверждения.\n"
        f"Используй /help для списка команд."
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отправляет список доступных команд."""
    username = update.message.from_user.username
    if username is None:
        return await update.message.reply_text("Пожалуйста, установите @username в настройках Telegram, чтобы использовать этот бот.")
    text = "Команды для всех:\n/start\n/help\n+ / -\nГеолокация"
    if is_leader(username):
        text += "\n\nКоманды для лидера:\n/report - сгенерировать отчёт\n/add_student <username> <ФИО> <role> - добавить пользователя\n/list_students - посмотреть список"
    elif is_student(username):
        text += "\n\nКоманды для студента:\n/status - узнать свой статус"
    await update.message.reply_text(text)

async def mark_attendance(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обрабатывает текстовые отметки '+' и '-'."""
    username = update.message.from_user.username
    if username is None:
        return await update.message.reply_text("Пожалуйста, установите @username.")
    
    text = update.message.text.strip()
    
    if not is_student(username):
        return await update.message.reply_text("Ты не в списке студентов.")
    
    today = datetime.now().strftime("%Y-%m-%d")
    if today not in attendance:
        attendance[today] = {}
        
    if text in ["+", "-"]:
        attendance[today][USERS[username]["name"]] = text
        await update.message.reply_text(f"Отметка сохранена: {text}")
    else:
        await update.message.reply_text("Используй только '+' или '-'")

async def location_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обрабатывает геолокацию пользователя."""
    username = update.message.from_user.username
    if username is None:
        return await update.message.reply_text("Пожалуйста, установите @username.")
    
    loc = update.message.location
    if not is_student(username):
        return await update.message.reply_text("Ты не в списке студентов.")
    
    user_coords = (loc.latitude, loc.longitude)
    dist = distance(user_coords, UNIVERSITY_CENTER)
    
    today = datetime.now().strftime("%Y-%m-%d")
    if today not in attendance:
        attendance[today] = {}
    
    if dist <= ALLOWED_RADIUS:
        attendance[today][USERS[username]["name"]] = "+"
        await update.message.reply_text("Ты в университете ✅")
    else:
        attendance[today][USERS[username]["name"]] = "-"
        await update.message.reply_text("Ты не в университете ❌")

# ================== ЛИДЕРСКИЕ КОМАНДЫ ==================

async def report(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Генерирует отчёт о посещаемости за текущий день в формате .docx."""
    username = update.message.from_user.username
    if not is_leader(username):
        return await update.message.reply_text("Только лидер может запрашивать отчёт.")
        
    today = datetime.now().strftime("%Y-%m-%d")
    if today not in attendance:
        return await update.message.reply_text("Сегодня ещё нет данных.")
        
    doc = Document()
    doc.add_heading(f"Отчёт за {today}", 0)
    for user in USERS.values():
        if user["role"] == "student":
            status = attendance[today].get(user["name"], "Не отмечен")
            doc.add_paragraph(f"{user['name']}: {status}")
            
    filename = f"attendance_{today}.docx"
    doc.save(filename)
    
    await update.message.reply_document(open(filename, "rb"))
    os.remove(filename)

async def add_student(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Добавляет нового пользователя (студента или лидера)."""
    username = update.message.from_user.username
    # Исправленная логика: разрешаем добавлять первого пользователя, если USERS пуст
    if not is_leader(username) and len(USERS) > 0:
        return await update.message.reply_text("Только лидер может добавлять пользователей.")
    
    try:
        args = context.args
        if len(args) < 3:
            return await update.message.reply_text("Используй: /add_student <username> <ФИО> <role>")
            
        new_username = args[0]
        # Собираем ФИО, если оно состоит из нескольких слов
        name = " ".join(args[1:-1]) 
        role = args[-1].lower() # Преобразуем роль в нижний регистр для единообразия
        
        if role not in ["student", "leader"]:
            return await update.message.reply_text("Роль должна быть student или leader")
            
        USERS[new_username] = {"name": name, "role": role}
        await update.message.reply_text(f"{name} ({role}) добавлен!")
        
    except Exception as e:
        await update.message.reply_text(f"Ошибка: {e}")

async def list_students(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Выводит список всех зарегистрированных пользователей."""
    username = update.message.from_user.username
    if not is_leader(username):
        return await update.message.reply_text("Только лидер может видеть список.")
        
    if not USERS:
        return await update.message.reply_text("Список пользователей пуст.")
        
    text = "Список пользователей:\n"
    for u, data in USERS.items():
        text += f"@{u}: {data['name']} ({data['role']})\n"
        
    await update.message.reply_text(text)

async def status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Показывает студенту его текущий статус посещаемости."""
    username = update.message.from_user.username
    if not is_student(username):
        return await update.message.reply_text("Ты не в списке студентов.")
        
    today = datetime.now().strftime("%Y-%m-%d")
    name = USERS[username]["name"]
    status_today = attendance.get(today, {}).get(name, "Не отмечен")
    
    await update.message.reply_text(f"{name}, твой статус сегодня: {status_today}")

# ================== ЗАПУСК БОТА ==================
def main():
    """Основная функция для запуска бота."""
    if TOKEN is None:
        logging.error("Telegram bot token not found. Set the 'TOKEN' environment variable.")
        return
    
    load_data() # <-- Добавьте эту строку, чтобы загрузить данные при запуске
    
    app = Application.builder().token(TOKEN).build()

    # Обработчики команд
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("report", report))
    app.add_handler(CommandHandler("add_student", add_student))
    app.add_handler(CommandHandler("list_students", list_students))
    app.add_handler(CommandHandler("status", status))

    # Обработчики сообщений
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, mark_attendance))
    app.add_handler(MessageHandler(filters.LOCATION, location_handler))

    logging.info("Бот запущен. Ожидание команд...")
    app.run_polling()

if __name__ == "__main__":
    main()
